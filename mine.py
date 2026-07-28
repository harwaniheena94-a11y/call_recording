

import os
from google import genai
from docx import Document
from dotenv import load_dotenv

# ==========================
# Load API Key
# ==========================

load_dotenv()

client = genai.Client(
    api_key=os.getenv("GEMINI_API_KEY")
)

# ==========================
# Load Gemini Model
# ==========================



# ==========================
# Read Transcript
# ==========================


with open(
    "transcripts/Donald Victor Williams-1.txt",
    "r",
    encoding="utf-8"
) as f:
    transcript = f.read()

# ==========================
# Prompt
# ==========================

PROMPT = f"""
chloe looks towards me from driving seat 
Viv , what's the plan , how we r going to celebrate your  today's victory

I continued scrolling through my phone and said 
"by sleeping and eating hehehe"

she groaned "you sloth , I knew , you had nothing planned that's why i bought vip passes of the new club"

"please exclude me , i am not going anywhere " i stated

"ohh sweetheart , i didn't even asked you , we are going that's final "
"emmm ,fine let's go home to change clothes "
chloe started laughing sracstically 
i looked at you weirdly , is she gone mad "what !"
"ohh what do you think i don't know you , and your tricks , i can smell your intentions babe "

i wased surprised , how did she knew about my plan "what, did you got electric shock or bloody merry hit you or something "

"bitch , you never agree to anything this easily ,i am damn sure once we go home , you will lock yourself in room , or possess me with some
 power so that i also stay at home , so  that's why we are not going home , we will buy new clothes and change their only  "

"haaw , i didn't even thought anything like that and i don't want to buy new dress , i am broke "

"bitch , atleast have some shame , if you forgot i am your manager also , i have all your bank info "

"ohhh",i lean in my seat disappointedly "how much time will it take "

"viv"

"i hate you bitch "i rolled my eyes

"i hate you too bitch" 



chater 2 

i hate to admit this but i really like shopping , chloe and i bought more than 5 new dresses and heels ,
"if i remember correctly , someone wants to eat and sleep only , but now she is enjoying shopping " chloe said sarcastically while i was 
admiring the heels , on which i got my eyes , "this is last one , after than we will change and go to club " i said
 oh my god i love heels ,and how could i even resist myself when  jimmy choe and christian louboutin had their new collection out  

"" 

after 2 hours of shopping , we finally reached the club , and i was amazed by the club's interior , it was so beautiful and luxurious
i was wearing a silk violet  backless dress with a slit on the left side , and a pair of silver shimmering heels , 
my blonde hair was open and i was wearing a diamond necklace and earrings and a tiara on my head ,ofcourse i love wearing tiara
, i think people should normalize wearing tiara , it looks so beautiful and elegant
, and chloe was wearing a dark cherry dress with a slit on the right side and a pair of black heels

we sat in our booth and ordered our drinks , i was sipping on my cocktail while chloe was enjoying her wine , 
we were enjoying the music and the vibe of the club ,one thing about chloe , she is a walking , alive friendship application , she
has a lot of friends and that friends of her has a lot of different friends and this girl know all of them , i just can't understand 
how she remembers who's friend is who and what's their name , because some have same name too , i just wish 
i never encounter her any friends whenver i am alone , because i have terrible memory even hamster has better memory than me , 
while i was lost in my thoughts , chloe invited a group of people in our booth 

and one thing about me , i don't like people , especially strangers , and even more when they are noisy and some of them are scary 
not some of them actually just one of them , he was a tall guy with a beard and a tattoo on his neck ,he looks like a hot male model 
who is out of straight vogue magazine , but out of budget ,he was wearing a black shirt with rolled up sleeves and a pair of black pants , even stone might have some expression on their 
 face but this guy's face was like a stone , i didn't dare to look at him 

generally i am quite fearless person , but this men aura is yelling trouble and danger , it better to stay away for the sake of peace in my life , 

there were two couch in the booth , placed opposite to each other , on one couch me , chloe and her friend natalia was sitting on another
couch there were 3 men , natalia boyfriend josh , and her brother jermey , and jeremy friend , the hot male model asher 

so in short chloe only knows natalia ,and only heard about her boyfriend josh that's it , but now here we all are sitting  together who fucking
meet each other for the first time , i can say from this mens expression this is not what they had planned for the night ,

after few drinks , everybody got quite loosen up , we dance , (we inclueds me , chloe , natalia and josh) , jermey and asher were 
sitting on couch and talking about something God knows what , but it was looking like they were planning to kill someone or something illegal

after dancing for a while , we were feeling like quite sober so we went to our booth to drink again , while we were drinking the waiter arrive and
gave me a note , natalia was talking to josh she stopped and looked at me quite excitedly like i got a diamond ring , 
i asked waiter who gave the note , he pointed towards the group of men sitting accros our booth , one of them was holding a cash bundle in his 
hand just to flex , i can see the disgust on natalia face all her excitment vanish when she saw that men , everybody was equally disgusted 
except asher , he was looking normal same poker face , who knows what he was thinking , i opened the note and read it , it says 
"spend the night with me  and i will give you as much money as you want "

chloe asked me what was written in the note , i told her , she was shocked and disgusted , "does he even know who you are 
, like how the fuck , let me teach that asshole a lesson " chole stand up , i 

"""

# ==========================
# Generate Report
# ==========================

response = client.models.generate_content(
    model="gemini-3.5-flash",
    contents=PROMPT
)

report = response.text

# ==========================
# Save Markdown
# ==========================

os.makedirs("reports", exist_ok=True)

with open(
    "reports/Donald Victor Williams-1.md",
    "w",
    encoding="utf-8"
) as f:
    f.write(report)

# ==========================
# Save Word Document
# ==========================

doc = Document()

doc.add_heading("Sales Call Analysis Report", level=1)

doc.add_paragraph(report)

doc.save("reports/Brian Behrmann.docx")

print("Report generated successfully!")
