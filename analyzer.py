import os
from google import genai
from docx import Document
from dotenv import load_dotenv

# ==========================
# Load API Key
# ==========================

load_dotenv()

client = genai.Client(
    api_key=os.getenv("GEMINI_API_KEY")
)

# ==========================
# Load Gemini Model
# ==========================



# ==========================
# Read Transcript
# ==========================
        

with open(
    "transcripts/Ninos Israel-4.txt",
    "r",
    encoding="utf-8"
) as f:
    transcript = f.read()

# ==========================
# Prompt
# ==========================

PROMPT = f"""
You are an experienced Sales Knowledge Analyst.

Analyze the following sales call transcript.

Generate a clean, concise, and professional Markdown report.

Keep the report under 500 words.

Do not invent information. If something is not mentioned, write "Not Mentioned."

Use the following structure:

#  Sales Call Knowledge Report

## Customer Overview
- Customer Name
- Location
- Products Discussed
- Call Outcome

## Call Summary
Summarize the conversation in 5–7 sentences.

## Customer Questions
Extract every meaningful question asked by the customer along with the answer provided by the salesperson.

## Customer Concerns
List the customer's concerns or objections as bullet points.

## Products & Features Discussed
List the products, services, rebates, financing options, or features discussed.

## Key Information Shared
List the most important information explained by the salesperson.

## Sales Tips Learned
Based on this conversation, write 3–5 practical lessons that would help a new sales representative.

## Potential FAQ
Generate 5–10 FAQ entries using only information from the conversation. Each FAQ should have:
- Question
- Answer

## Keywords
List the important keywords from the conversation.

Transcript:
{transcript}

"""

# ==========================
# Generate Report
# ==========================

response = client.models.generate_content(
    model="gemini-3.5-flash",
    contents=PROMPT
)

report = response.text

# ==========================
# Save Markdown
# ==========================

os.makedirs("reports", exist_ok=True)

with open(
    "reports/Ninos Israel-4.md",
    "w",
    encoding="utf-8"
) as f:
    f.write(report)

# ==========================
# Save Word Document
# ==========================

# doc = Document()

# doc.add_heading("Sales Call Analysis Report", level=1)

# doc.add_paragraph(report)

# doc.save("reports/Kelvin Redd.docx")

# print("Report generated successfully!")
